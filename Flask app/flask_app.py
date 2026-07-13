# -*- coding: utf-8 -*-
"""
Flask app: upload a CSV of links + a list of phrases, download a Word
document containing a QR code for every link with a random phrase
underneath it. Generation runs in a background thread so the page can
show a live progress bar while the QR codes are being created.

Run locally with:
    pip install -r requirements.txt
    python app.py
Then open http://127.0.0.1:5000

Note: progress is tracked in an in-memory dict, which only works with a
single-process server (fine for `python app.py` / `flask run`). If you
deploy behind multiple gunicorn workers, move `jobs` to something shared
like Redis instead.
"""

import csv
import io
import random
import threading
import time
import uuid
from datetime import datetime

# How long a finished/errored job is kept in memory after completion,
# in case the browser needs to retry the download. Swept on each /start.
JOB_RETENTION_SECONDS = 600

import qrcode
from flask import Flask, render_template, request, send_file, jsonify, abort

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE

app = Flask(__name__)
app.secret_key = "change-me-in-production"  # needed for session/flash use

# Shown as the pre-filled default text in the phrases textbox.
DEFAULT_PHRASES = """Keep up the good work!
You're doing amazing!
Fantastic effort!
Nice work!
Keep pushing forward!
You've got this!
Awesome progress!
Stay curious!
Well done!
Great job!
Terrific work!
Bazinga!
FTW!
You crushed it!
That’s how it’s done!
A masterclass in excellence!
Not even gravity can bring you down!
You’re on fire! (Not literally, I hope)
The world needs more people like you!
You make this look easy!
Standing ovation from me!
Keep being awesome!
10/10, would recommend!
You deserve all the cake!
Absolute legend!
Someone call the hall of fame!
You + effort = magic!
A true force of nature!
You didn’t just raise the bar—you launched it into orbit!
That was smoother than a jazz solo!
You're like Wi-Fi: consistently strong!
Whatever you just did? Do it again!
Scientists are studying how you’re this good!
I feel like I should be taking notes!
You're officially too good to quit!
You understood the assignment!
Excellence runs in your veins!
Give yourself a high-five!
Your talent deserves a parade!
Your keyboard must be smoking from that brilliance!
A standing ovation—but in my digital world!
You’ve mastered this more than a Jedi masters the Force!
Is this your superhero origin story?
Big brain energy at its finest!
You’re the human version of a five-star review!
Brain = galaxy-level genius!
I'd frame this moment if I could!
If there were an award for this, you'd win it!
Like a pro athlete, but for pure excellence!
The student becomes the master!
*Chef’s kiss*—perfect execution!
Consider the mic dropped!
You should teach a masterclass in this!
If success had a sound, it’d be your footsteps!
Is this greatness? Yes, yes it is.
You didn’t just do well—you did **phenomenally**!
Even the Internet applauds you!
Someone hand you a trophy already!"""

# In-memory job store: job_id -> dict with progress/result info.
jobs = {}
jobs_lock = threading.Lock()


def build_qr_document(dist_list, phrases, progress_callback=None):
    """Given a list of CSV row dicts and a list of phrase strings,
    build and return an in-memory .docx Document.

    progress_callback(done, total) is called after each QR code is
    placed, if provided.
    """

    # Find the "Link" column; fall back to the first column whose
    # values look like a URL.
    link_col = "Link"
    if dist_list and link_col not in dist_list[0]:
        for col in dist_list[0]:
            value = dist_list[0][col]
            if isinstance(value, str) and value.strip().lower().startswith("http"):
                link_col = col
                break

    document = Document()

    # 1.5cm margins
    for section in document.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    if not dist_list:
        document.add_paragraph("No links were found in the uploaded CSV file.")
        return document

    n_rows = (len(dist_list) // 3) + (1 if len(dist_list) % 3 else 0)
    table = document.add_table(rows=n_rows, cols=3, style="Table Grid")

    for row in table.rows:
        row.height = Cm(4.8)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    total = len(dist_list)
    for i, row in enumerate(dist_list):
        link_value = row.get(link_col, "")
        if link_value:
            img = qrcode.make(link_value)
            image_stream = io.BytesIO()
            img.save(image_stream)
            image_stream.seek(0)

            cell = table.cell(i // 3, i % 3)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = cell.paragraphs[0].add_run()
            run.add_picture(image_stream, width=Cm(3.8), height=Cm(3.8))

            phrase = random.choice(phrases) if phrases else ""
            cell.paragraphs[0].add_run("\n" + phrase)

        if progress_callback:
            progress_callback(i + 1, total)

    return document


def parse_csv(file_bytes):
    """Read raw CSV bytes into a list of dicts."""
    stream = io.StringIO(file_bytes.decode("utf-8-sig"))
    reader = csv.DictReader(stream)
    return list(reader)


def parse_phrases(raw_text):
    """Turn the textarea contents into a clean list of non-empty phrases."""
    lines = [line.strip() for line in raw_text.splitlines()]
    return [line for line in lines if line]


def run_job(job_id, dist_list, phrases):
    """Background worker: builds the document and updates job progress."""

    def on_progress(done, total):
        with jobs_lock:
            jobs[job_id]["current"] = done
            jobs[job_id]["total"] = total

    try:
        document = build_qr_document(dist_list, phrases, progress_callback=on_progress)
        output_stream = io.BytesIO()
        document.save(output_stream)
        output_stream.seek(0)

        timestamp = datetime.now().strftime("%Y-%m-%d_%H%M")
        with jobs_lock:
            jobs[job_id]["status"] = "done"
            # Store raw bytes, not the BytesIO object itself: send_file()
            # closes/consumes whatever stream it's given, so a shared
            # stream object would break a second download of the same job.
            jobs[job_id]["file_bytes"] = output_stream.getvalue()
            jobs[job_id]["filename"] = f"QR Codes {timestamp}.docx"
            jobs[job_id]["finished_at"] = time.time()
    except Exception as exc:  # noqa: BLE001
        with jobs_lock:
            jobs[job_id]["status"] = "error"
            jobs[job_id]["error"] = str(exc)
            jobs[job_id]["finished_at"] = time.time()


def sweep_old_jobs():
    """Remove finished/errored jobs older than JOB_RETENTION_SECONDS.
    Called opportunistically from /start so no separate scheduler is needed."""
    cutoff = time.time() - JOB_RETENTION_SECONDS
    with jobs_lock:
        stale_ids = [
            jid
            for jid, job in jobs.items()
            if job["status"] in ("done", "error") and job.get("finished_at", 0) < cutoff
        ]
        for jid in stale_ids:
            jobs.pop(jid, None)


@app.route("/", methods=["GET"])
def index():
    return render_template("index.html", default_phrases=DEFAULT_PHRASES)


@app.route("/start", methods=["POST"])
def start():
    """Validate the upload, kick off a background job, and return its id."""
    sweep_old_jobs()

    csv_file = request.files.get("csv_file")
    if not csv_file or csv_file.filename == "":
        return jsonify({"error": "Please upload a CSV file."}), 400

    if not csv_file.filename.lower().endswith(".csv"):
        return jsonify({"error": "The uploaded file must be a .csv file."}), 400

    phrases_raw = request.form.get("phrases", "")
    phrases = parse_phrases(phrases_raw) or parse_phrases(DEFAULT_PHRASES)

    try:
        dist_list = parse_csv(csv_file.read())
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": f"Could not read the CSV file: {exc}"}), 400

    if not dist_list:
        return jsonify({"error": "The CSV file appears to be empty."}), 400

    job_id = uuid.uuid4().hex
    with jobs_lock:
        jobs[job_id] = {
            "status": "running",
            "current": 0,
            "total": len(dist_list),
            "file_bytes": None,
            "filename": None,
            "error": None,
        }

    thread = threading.Thread(target=run_job, args=(job_id, dist_list, phrases), daemon=True)
    thread.start()

    return jsonify({"job_id": job_id})


@app.route("/progress/<job_id>", methods=["GET"])
def progress(job_id):
    with jobs_lock:
        job = jobs.get(job_id)
        if not job:
            abort(404)
        return jsonify(
            {
                "status": job["status"],
                "current": job["current"],
                "total": job["total"],
                "error": job["error"],
            }
        )


@app.route("/download/<job_id>", methods=["GET"])
def download(job_id):
    with jobs_lock:
        job = jobs.get(job_id)
        if not job or job["status"] != "done":
            abort(404)
        # Fresh BytesIO each time: send_file() closes whatever stream
        # it's given, so a shared stream object would break a repeat
        # download of the same job.
        file_stream = io.BytesIO(job["file_bytes"])
        filename = job["filename"]

    # Note: we deliberately do NOT delete the job here. If the browser's
    # progress poll fires twice near completion (setInterval can overlap
    # if a fetch is slow), both requests may try to download the same
    # job; deleting it on first download would 404 the second one. Stale
    # jobs are instead cleaned up later by sweep_old_jobs().
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    app.run(debug=True, threaded=True)
